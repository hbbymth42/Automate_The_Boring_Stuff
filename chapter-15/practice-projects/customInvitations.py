#! python3
# customInvitations.py - Create Custom Invitations from text file Guest List

import docx
from pathlib import Path

guestFile = open(Path.cwd() / 'guests.txt')

guestList = guestFile.readlines()

doc = docx.Document('Invitations.docx')

iterator = 0

for guestNum, guest in enumerate(guestList):
    if guestNum == 0:
        doc.paragraphs[0].text = 'It would be a pleasure to have the company of'
        doc.paragraphs[0].style = 'Invite_Text'
        iterator = iterator + 1
        doc.add_paragraph(guest)
        doc.paragraphs[iterator].style = 'Invite_Name'
        iterator = iterator + 1
        doc.add_paragraph('at')
        doc.paragraphs[iterator].add_run(' 11010 Memory Lane on the Evening of')
        doc.paragraphs[iterator].style = 'Invite_Text'
        doc.paragraphs[iterator].runs[0].underline = True
        iterator = iterator + 1
        doc.add_paragraph('April 1st')
        doc.paragraphs[iterator].style = 'Invite_Date'
        iterator = iterator + 1
        doc.add_paragraph('at')
        doc.paragraphs[iterator].add_run(" 7 o'clock")
        doc.paragraphs[iterator].style = 'Invite_Text'
        doc.paragraphs[iterator].runs[0].underline = True
        doc.paragraphs[iterator].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
        iterator = iterator + 1
    else:
        doc.add_paragraph('It would be a pleasure to have the company of')
        doc.paragraphs[iterator].style = 'Invite_Text'
        iterator = iterator + 1
        doc.add_paragraph(guest)
        doc.paragraphs[iterator].style = 'Invite_Name'
        iterator = iterator + 1
        doc.add_paragraph('at')
        doc.paragraphs[iterator].add_run(' 11010 Memory Lane on the Evening of')
        doc.paragraphs[iterator].style = 'Invite_Text'
        doc.paragraphs[iterator].runs[0].underline = True
        iterator = iterator + 1
        doc.add_paragraph('April 1st')
        doc.paragraphs[iterator].style = 'Invite_Date'
        iterator = iterator + 1
        doc.add_paragraph('at')
        doc.paragraphs[iterator].add_run(" 7 o'clock")
        doc.paragraphs[iterator].style = 'Invite_Text'
        doc.paragraphs[iterator].runs[0].underline = True
        doc.paragraphs[iterator].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
        iterator = iterator + 1

doc.save('Invitations.docx')